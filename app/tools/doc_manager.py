import os
import json
from pathlib import Path
from typing import Dict, Any, List, Optional
from pypdf import PdfReader
import docx
from app.config import settings
from app.utils.logger import app_logger, audit_logger
from app.policy import PolicyEvaluator

class DocumentManager:
    WORKSPACE_DIR = settings.DATA_DIR / "workspace"
    APPROVED_DOCS_DIR = settings.DATA_DIR / "approved_docs"

    @classmethod
    def ensure_directories(cls):
        cls.WORKSPACE_DIR.mkdir(parents=True, exist_ok=True)
        cls.APPROVED_DOCS_DIR.mkdir(parents=True, exist_ok=True)

    @classmethod
    def _resolve_file_path(cls, file_path_str: str) -> Path:
        cls.ensure_directories()
        path = Path(file_path_str)
        if not path.is_absolute():
            # Default relative paths to data/workspace if not prefixed
            path = settings.BASE_DIR / path
        return path

    @classmethod
    def read_document(cls, file_path_str: str) -> Dict[str, Any]:
        """
        Reads .pdf, .docx, .txt, .md, .json, .py, .csv, .log files and returns content.
        """
        file_path = cls._resolve_file_path(file_path_str)

        if not file_path.exists():
            return {
                "success": False,
                "error": f"File not found: '{file_path}'",
                "file_name": file_path.name,
                "content": ""
            }

        ext = file_path.suffix.lower()
        extracted_text = ""

        try:
            if ext == ".pdf":
                reader = PdfReader(file_path)
                pages_text = [page.extract_text() for page in reader.pages if page.extract_text()]
                extracted_text = "\n\n".join(pages_text)

            elif ext == ".docx":
                doc = docx.Document(file_path)
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                extracted_text = "\n\n".join(paragraphs)

            elif ext in [".txt", ".md", ".json", ".py", ".csv", ".log", ".yaml", ".yml", ".html", ".css", ".js"]:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    extracted_text = f.read()

            else:
                return {
                    "success": False,
                    "error": f"Unsupported file format for reading: '{ext}'",
                    "file_name": file_path.name,
                    "content": ""
                }

            return {
                "success": True,
                "file_name": file_path.name,
                "file_path": str(file_path),
                "extension": ext,
                "content": extracted_text[:20000],  # Truncate to safe character limit
                "length": len(extracted_text)
            }
        except Exception as e:
            app_logger.error(f"Error reading document '{file_path}': {e}")
            return {
                "success": False,
                "error": f"Error reading document: {str(e)}",
                "file_name": file_path.name,
                "content": ""
            }

    @classmethod
    def create_document(cls, file_path_str: str, content: str, overwrite: bool = False) -> Dict[str, Any]:
        """
        Creates a new document (.txt, .md, .json, .py, .docx) in approved workspace.
        """
        file_path = cls._resolve_file_path(file_path_str)

        # Check safety policy (Level 1: Draft / Write in designated folders)
        allowed, reason, level = PolicyEvaluator.evaluate_action("write_draft", {"file_path": str(file_path)})
        if not allowed:
            return {"success": False, "error": f"Policy Blocked: {reason}", "level": level}

        if file_path.exists() and not overwrite:
            return {
                "success": False,
                "error": f"File '{file_path.name}' already exists. Pass overwrite=True to replace.",
                "file_path": str(file_path)
            }

        ext = file_path.suffix.lower()

        try:
            file_path.parent.mkdir(parents=True, exist_ok=True)

            if ext == ".docx":
                doc = docx.Document()
                for line in content.split("\n"):
                    doc.add_paragraph(line)
                doc.save(file_path)
            else:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(content)

            audit_logger.info(f"Created document '{file_path.name}' at {file_path}")
            return {
                "success": True,
                "message": f"Successfully created document '{file_path.name}'.",
                "file_name": file_path.name,
                "file_path": str(file_path),
                "size_bytes": file_path.stat().st_size
            }
        except Exception as e:
            app_logger.error(f"Error creating document '{file_path}': {e}")
            return {"success": False, "error": f"Failed to create document: {str(e)}"}

    @classmethod
    def edit_document(
        cls, 
        file_path_str: str, 
        new_content: Optional[str] = None, 
        append_content: Optional[str] = None,
        search_target: Optional[str] = None,
        replace_text: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Edits an existing document by replacing full content, appending text, or search-and-replacing text.
        """
        file_path = cls._resolve_file_path(file_path_str)

        if not file_path.exists():
            return {"success": False, "error": f"File not found for editing: '{file_path}'"}

        ext = file_path.suffix.lower()
        if ext == ".pdf":
            return {"success": False, "error": "Editing PDF files directly is not supported. Create a new DOCX or MD file instead."}

        try:
            if new_content is not None:
                # Full overwrite
                return cls.create_document(str(file_path), new_content, overwrite=True)

            elif append_content is not None:
                # Append text
                if ext == ".docx":
                    doc = docx.Document(file_path)
                    for line in append_content.split("\n"):
                        doc.add_paragraph(line)
                    doc.save(file_path)
                else:
                    with open(file_path, "a", encoding="utf-8") as f:
                        f.write(f"\n{append_content}")
                return {
                    "success": True,
                    "message": f"Appended text to '{file_path.name}'.",
                    "file_path": str(file_path)
                }

            elif search_target is not None and replace_text is not None:
                # Search and replace
                read_res = cls.read_document(str(file_path))
                if not read_res["success"]:
                    return read_res
                
                updated_text = read_res["content"].replace(search_target, replace_text)
                return cls.create_document(str(file_path), updated_text, overwrite=True)

            return {"success": False, "error": "No edit instruction provided (pass new_content, append_content, or search_target + replace_text)."}

        except Exception as e:
            app_logger.error(f"Error editing document '{file_path}': {e}")
            return {"success": False, "error": f"Failed to edit document: {str(e)}"}

    @classmethod
    def list_workspace_files(cls) -> List[Dict[str, Any]]:
        """
        Lists all files in data/workspace/ and data/approved_docs/ directories.
        """
        cls.ensure_directories()
        files = []
        for folder in [cls.WORKSPACE_DIR, cls.APPROVED_DOCS_DIR]:
            for root, _, filenames in os.walk(folder):
                for filename in filenames:
                    p = Path(root) / filename
                    files.append({
                        "file_name": p.name,
                        "file_path": str(p),
                        "relative_path": str(p.relative_to(settings.BASE_DIR)),
                        "size_bytes": p.stat().st_size,
                        "extension": p.suffix.lower()
                    })
        return files
